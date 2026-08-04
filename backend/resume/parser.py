"""
Resume Analysis Engine - Phase 3.2
Intelligent extraction and structuring of resume information from PDF/DOCX files
"""
import re
import logging
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from pathlib import Path
import PyPDF2
from docx import Document
import email_validator
from email_validator import validate_email
from urllib.parse import urlparse

logger = logging.getLogger(__name__)


@dataclass
class ExtractedResumeData:
    """Structured resume data extracted by the Resume Analysis Engine"""
    # Personal Information
    full_name: str = ""
    email: str = ""
    phone: str = ""
    linkedin_url: str = ""
    github_url: str = ""
    portfolio_url: str = ""
    location: str = ""
    
    # Professional Summary
    professional_summary: str = ""
    
    # Skills
    technical_skills: List[str] = None
    soft_skills: List[str] = None
    programming_languages: List[str] = None
    frameworks: List[str] = None
    databases: List[str] = None
    cloud_platforms: List[str] = None
    tools: List[str] = None
    
    # Work Experience
    work_experience: List[Dict[str, Any]] = None
    
    # Projects
    projects: List[Dict[str, Any]] = None
    
    # Education
    education: List[Dict[str, Any]] = None
    
    # Additional Information
    certifications: List[str] = None
    languages: List[str] = None
    achievements: List[str] = None
    publications: List[str] = None
    volunteer_experience: List[Dict[str, Any]] = None
    
    # Metadata
    sections_found: List[str] = None
    extraction_confidence: float = 0.0
    processing_notes: Dict[str, Any] = None
    
    def __post_init__(self):
        """Initialize lists to avoid mutable defaults"""
        if self.technical_skills is None:
            self.technical_skills = []
        if self.soft_skills is None:
            self.soft_skills = []
        if self.programming_languages is None:
            self.programming_languages = []
        if self.frameworks is None:
            self.frameworks = []
        if self.databases is None:
            self.databases = []
        if self.cloud_platforms is None:
            self.cloud_platforms = []
        if self.tools is None:
            self.tools = []
        if self.work_experience is None:
            self.work_experience = []
        if self.projects is None:
            self.projects = []
        if self.education is None:
            self.education = []
        if self.certifications is None:
            self.certifications = []
        if self.languages is None:
            self.languages = []
        if self.achievements is None:
            self.achievements = []
        if self.publications is None:
            self.publications = []
        if self.volunteer_experience is None:
            self.volunteer_experience = []
        if self.sections_found is None:
            self.sections_found = []
        if self.processing_notes is None:
            self.processing_notes = {}


class ResumeParsingError(Exception):
    """Custom exception for resume parsing errors"""
    pass


class ResumeAnalysisEngine:
    """
    Intelligent Resume Analysis Engine for Phase 3.2
    Extracts structured information from PDF/DOCX resume files
    """
    
    def __init__(self):
        """Initialize the Resume Analysis Engine"""
        self.technical_skills_db = self._load_technical_skills()
        self.soft_skills_db = self._load_soft_skills()
        self.section_patterns = self._load_section_patterns()
        
    def analyze_resume_file(self, file_path: str, file_type: str) -> ExtractedResumeData:
        """
        Main entry point for resume analysis
        
        Args:
            file_path: Path to the resume file
            file_type: Type of file ('pdf' or 'docx')
            
        Returns:
            ExtractedResumeData with structured information
            
        Raises:
            ResumeParsingError: If parsing fails
        """
        try:
            # Extract raw text from file
            raw_text = self._extract_text_from_file(file_path, file_type)
            
            # Analyze the extracted text
            return self.analyze_resume_text(raw_text)
            
        except Exception as e:
            logger.error(f"Resume analysis failed for {file_path}: {e}")
            raise ResumeParsingError(f"Failed to analyze resume: {str(e)}")
    
    def analyze_resume_text(self, resume_text: str) -> ExtractedResumeData:
        """
        Analyze resume text and extract structured information
        
        Args:
            resume_text: Raw text extracted from resume
            
        Returns:
            ExtractedResumeData with structured information
        """
        if not resume_text or len(resume_text.strip()) < 10:
            raise ResumeParsingError("Resume text is empty or too short")
        
        logger.info("Starting resume text analysis...")
        
        # Initialize result data
        result = ExtractedResumeData()
        
        # Clean and normalize text
        clean_text = self._clean_text(resume_text)
        
        # Detect resume sections
        sections = self._detect_sections(clean_text)
        result.sections_found = list(sections.keys())
        
        # Extract personal information
        self._extract_personal_info(clean_text, result)
        
        # Extract professional summary
        result.professional_summary = self._extract_professional_summary(sections, clean_text)
        
        # Extract skills
        self._extract_skills(sections, clean_text, result)
        
        # Extract work experience
        result.work_experience = self._extract_work_experience(sections, clean_text)
        
        # Extract projects
        result.projects = self._extract_projects(sections, clean_text)
        
        # Extract education
        result.education = self._extract_education(sections, clean_text)
        
        # Extract additional information
        result.certifications = self._extract_certifications(sections, clean_text)
        result.languages = self._extract_languages(sections, clean_text)
        result.achievements = self._extract_achievements(sections, clean_text)
        result.publications = self._extract_publications(sections, clean_text)
        result.volunteer_experience = self._extract_volunteer_experience(sections, clean_text)
        
        # Calculate extraction confidence
        result.extraction_confidence = self._calculate_confidence(result)
        
        # Add processing notes
        result.processing_notes = {
            'total_sections_detected': len(result.sections_found),
            'has_contact_info': bool(result.full_name and result.email),
            'has_work_experience': len(result.work_experience) > 0,
            'total_skills_found': len(result.technical_skills) + len(result.soft_skills)
        }
        
        logger.info(f"Resume analysis completed. Confidence: {result.extraction_confidence:.2f}")
        return result
    
    def _extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extract raw text from PDF or DOCX file"""
        try:
            if file_type.lower() == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_type.lower() == 'docx':
                return self._extract_docx_text(file_path)
            else:
                raise ResumeParsingError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            raise ResumeParsingError(f"Could not extract text from {file_type} file")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file using PyPDF2"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            if not text.strip():
                raise ResumeParsingError("No text found in PDF file")
                
            return text
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise ResumeParsingError("Failed to read PDF file")
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file using python-docx"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise ResumeParsingError("No text found in DOCX file")
                
            return text
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise ResumeParsingError("Failed to read DOCX file")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize resume text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common encoding issues
        text = text.replace('•', '•').replace('–', '-').replace('"', '"').replace('"', '"')
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def _load_section_patterns(self) -> Dict[str, List[str]]:
        """Load regex patterns for detecting resume sections"""
        return {
            'contact': [
                r'contact\s*(information|info)?',
                r'personal\s*(information|info|details)',
                r'profile',
            ],
            'summary': [
                r'professional\s*summary',
                r'career\s*(summary|objective)',
                r'summary\s*(of\s*qualifications)?',
                r'objective',
                r'profile',
                r'about\s*(me)?',
            ],
            'skills': [
                r'technical\s*skills',
                r'core\s*skills',
                r'skills\s*(and\s*competencies)?',
                r'competencies',
                r'expertise',
                r'proficiencies',
            ],
            'experience': [
                r'professional\s*experience',
                r'work\s*experience',
                r'employment\s*(history)?',
                r'career\s*history',
                r'experience',
            ],
            'projects': [
                r'projects?',
                r'academic\s*projects?',
                r'personal\s*projects?',
                r'portfolio',
                r'key\s*projects?',
            ],
            'education': [
                r'education',
                r'academic\s*(background|qualifications)',
                r'qualifications',
                r'degrees?',
            ],
            'certifications': [
                r'certifications?',
                r'certificates?',
                r'professional\s*certifications?',
                r'licenses?',
            ],
            'languages': [
                r'languages?',
                r'foreign\s*languages?',
                r'programming\s*languages?',
            ],
            'achievements': [
                r'achievements?',
                r'awards?',
                r'honors?',
                r'accomplishments?',
                r'recognition',
            ],
            'publications': [
                r'publications?',
                r'research\s*papers?',
                r'papers?',
            ],
            'volunteer': [
                r'volunteer\s*(experience|work)',
                r'community\s*(service|involvement)',
                r'volunteering',
            ],
        }
    
    def _load_technical_skills(self) -> Dict[str, List[str]]:
        """Load comprehensive technical skills database"""
        return {
            'programming_languages': [
                'Python', 'JavaScript', 'Java', 'C++', 'C#', 'C', 'PHP', 'Ruby', 'Go', 'Rust',
                'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB', 'TypeScript', 'Dart', 'Perl',
                'Shell', 'Bash', 'PowerShell', 'SQL', 'HTML', 'CSS', 'SCSS', 'SASS',
            ],
            'frameworks': [
                'React', 'Vue.js', 'Angular', 'Django', 'Flask', 'Spring', 'Express.js',
                'Node.js', 'Laravel', 'Ruby on Rails', 'ASP.NET', '.NET', 'Flutter',
                'React Native', 'Xamarin', 'Bootstrap', 'Tailwind CSS', 'jQuery',
                'Next.js', 'Nuxt.js', 'Svelte', 'Ember.js', 'Backbone.js',
            ],
            'databases': [
                'MySQL', 'PostgreSQL', 'MongoDB', 'SQLite', 'Oracle', 'SQL Server',
                'Redis', 'Cassandra', 'DynamoDB', 'Elasticsearch', 'Neo4j',
                'CouchDB', 'Firebase', 'Supabase', 'MariaDB', 'InfluxDB',
            ],
            'cloud_platforms': [
                'AWS', 'Azure', 'Google Cloud', 'GCP', 'Heroku', 'DigitalOcean',
                'Vercel', 'Netlify', 'Firebase', 'Cloudflare', 'IBM Cloud',
                'Oracle Cloud', 'Alibaba Cloud',
            ],
            'tools': [
                'Git', 'GitHub', 'GitLab', 'Docker', 'Kubernetes', 'Jenkins',
                'Travis CI', 'CircleCI', 'Terraform', 'Ansible', 'Webpack',
                'Babel', 'ESLint', 'Prettier', 'Jest', 'Cypress', 'Selenium',
                'Postman', 'Swagger', 'VS Code', 'IntelliJ', 'Eclipse',
                'Vim', 'Emacs', 'Sublime Text', 'Figma', 'Sketch', 'Adobe XD',
            ],
            'methodologies': [
                'Agile', 'Scrum', 'Kanban', 'DevOps', 'CI/CD', 'TDD', 'BDD',
                'Microservices', 'RESTful APIs', 'GraphQL', 'SOAP', 'Serverless',
                'Machine Learning', 'Deep Learning', 'Data Science', 'Big Data',
                'Blockchain', 'IoT', 'Cybersecurity', 'Penetration Testing',
            ],
        }
    
    def _load_soft_skills(self) -> List[str]:
        """Load soft skills database"""
        return [
            'Leadership', 'Communication', 'Teamwork', 'Problem Solving',
            'Critical Thinking', 'Time Management', 'Project Management',
            'Analytical Thinking', 'Creativity', 'Adaptability', 'Flexibility',
            'Attention to Detail', 'Organizational Skills', 'Collaboration',
            'Mentoring', 'Training', 'Presentation Skills', 'Public Speaking',
            'Customer Service', 'Negotiation', 'Conflict Resolution',
            'Strategic Planning', 'Decision Making', 'Innovation',
            'Emotional Intelligence', 'Cultural Sensitivity', 'Multitasking',
        ]
    
    def _detect_sections(self, text: str) -> Dict[str, str]:
        """Detect and extract resume sections using intelligent pattern matching"""
        sections = {}
        text_lower = text.lower()
        lines = text.split('\n')
        
        current_section = None
        section_content = []
        
        for i, line in enumerate(lines):
            line_clean = line.strip()
            if not line_clean:
                continue
                
            # Check if this line is a section header
            detected_section = self._identify_section_header(line_clean)
            
            if detected_section:
                # Save previous section
                if current_section and section_content:
                    sections[current_section] = '\n'.join(section_content).strip()
                
                # Start new section
                current_section = detected_section
                section_content = []
            elif current_section:
                section_content.append(line_clean)
        
        # Save the last section
        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content).strip()
        
        return sections
    
    def _identify_section_header(self, line: str) -> Optional[str]:
        """Identify if a line is a section header"""
        line_lower = line.lower().strip()
        
        # Skip very short lines
        if len(line_lower) < 3:
            return None
        
        # Check against section patterns
        for section_type, patterns in self.section_patterns.items():
            for pattern in patterns:
                if re.match(rf'^{pattern}s?:?\s*$', line_lower):
                    return section_type
        
        return None
    
    def _extract_personal_info(self, text: str, result: ExtractedResumeData):
        """Extract personal contact information"""
        # Extract name (usually first line or after certain patterns)
        result.full_name = self._extract_name(text)
        
        # Extract email
        result.email = self._extract_email(text)
        
        # Extract phone
        result.phone = self._extract_phone(text)
        
        # Extract URLs
        result.linkedin_url = self._extract_linkedin_url(text)
        result.github_url = self._extract_github_url(text)
        result.portfolio_url = self._extract_portfolio_url(text)
        
        # Extract location
        result.location = self._extract_location(text)
    
    def _extract_name(self, text: str) -> str:
        """Extract full name from resume"""
        lines = text.split('\n')
        
        # Try first few lines
        for line in lines[:5]:
            line = line.strip()
            if len(line) > 5 and len(line) < 50:
                # Check if it looks like a name (2-4 words, mostly letters)
                words = line.split()
                if 2 <= len(words) <= 4:
                    if all(word.replace('.', '').replace(',', '').isalpha() and len(word) > 1 for word in words):
                        return line
        
        return ""
    
    def _extract_email(self, text: str) -> str:
        """Extract email address from resume"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, text)
        
        for email in matches:
            try:
                email_validator.validate_email(email)
                return email
            except:
                continue
        
        return ""
    
    def _extract_phone(self, text: str) -> str:
        """Extract phone number from resume"""
        # Various phone number patterns
        phone_patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\d{10}',
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            if matches:
                return matches[0]
        
        return ""
    
    def _extract_linkedin_url(self, text: str) -> str:
        """Extract LinkedIn profile URL"""
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9_-]+'
        matches = re.findall(linkedin_pattern, text, re.IGNORECASE)
        return matches[0] if matches else ""
    
    def _extract_github_url(self, text: str) -> str:
        """Extract GitHub profile URL"""
        github_pattern = r'(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9_-]+'
        matches = re.findall(github_pattern, text, re.IGNORECASE)
        return matches[0] if matches else ""
    
    def _extract_portfolio_url(self, text: str) -> str:
        """Extract portfolio website URL"""
        # Look for URLs that are not LinkedIn or GitHub
        url_pattern = r'https?://[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?:/[^\s]*)?'
        matches = re.findall(url_pattern, text, re.IGNORECASE)
        
        for url in matches:
            if 'linkedin.com' not in url.lower() and 'github.com' not in url.lower():
                return url
        
        return ""
    
    def _extract_location(self, text: str) -> str:
        """Extract location information"""
        # Look for city, state patterns or common location indicators
        location_patterns = [
            r'([A-Za-z\s]+),\s*([A-Z]{2})\s*\d{5}',  # City, State ZIP
            r'([A-Za-z\s]+),\s*([A-Z]{2})',          # City, State
            r'([A-Za-z\s]+),\s*([A-Za-z\s]+)',       # City, Country
        ]
        
        for pattern in location_patterns:
            matches = re.findall(pattern, text)
            if matches:
                return f"{matches[0][0]}, {matches[0][1]}"
        
        return ""
    
    def _extract_professional_summary(self, sections: Dict[str, str], text: str) -> str:
        """Extract professional summary or objective"""
        # Try to find in dedicated summary section
        for section_key in ['summary', 'objective', 'profile']:
            if section_key in sections:
                return sections[section_key]
        
        # Fallback: look for summary-like content in first few paragraphs
        lines = text.split('\n')
        for i, line in enumerate(lines[:10]):
            line = line.strip()
            if len(line) > 50 and any(keyword in line.lower() for keyword in 
                                    ['experienced', 'professional', 'skilled', 'passionate']):
                # Try to get the full paragraph
                paragraph = []
                for j in range(i, min(i + 5, len(lines))):
                    if lines[j].strip():
                        paragraph.append(lines[j].strip())
                    else:
                        break
                
                if len(' '.join(paragraph)) > 50:
                    return ' '.join(paragraph)
        
        return ""
    
    def _extract_skills(self, sections: Dict[str, str], text: str, result: ExtractedResumeData):
        """Extract and categorize skills"""
        # Combine skills sections
        skills_text = ""
        for section_key in ['skills', 'technical_skills', 'competencies']:
            if section_key in sections:
                skills_text += sections[section_key] + " "
        
        # If no dedicated skills section, search entire text
        if not skills_text.strip():
            skills_text = text
        
        # Extract technical skills by category
        for category, skills_list in self.technical_skills_db.items():
            found_skills = []
            for skill in skills_list:
                if self._skill_mentioned(skills_text, skill):
                    found_skills.append(skill)
            
            if category == 'programming_languages':
                result.programming_languages.extend(found_skills)
            elif category == 'frameworks':
                result.frameworks.extend(found_skills)
            elif category == 'databases':
                result.databases.extend(found_skills)
            elif category == 'cloud_platforms':
                result.cloud_platforms.extend(found_skills)
            elif category == 'tools':
                result.tools.extend(found_skills)
            else:
                result.technical_skills.extend(found_skills)
        
        # Extract soft skills
        for skill in self.soft_skills_db:
            if self._skill_mentioned(skills_text, skill):
                result.soft_skills.append(skill)
        
        # Remove duplicates while preserving order
        result.technical_skills = list(dict.fromkeys(result.technical_skills))
        result.soft_skills = list(dict.fromkeys(result.soft_skills))
        result.programming_languages = list(dict.fromkeys(result.programming_languages))
        result.frameworks = list(dict.fromkeys(result.frameworks))
        result.databases = list(dict.fromkeys(result.databases))
        result.cloud_platforms = list(dict.fromkeys(result.cloud_platforms))
        result.tools = list(dict.fromkeys(result.tools))
    
    def _skill_mentioned(self, text: str, skill: str) -> bool:
        """Check if a skill is mentioned in the text"""
        text_lower = text.lower()
        skill_lower = skill.lower()
        
        # Exact match with word boundaries
        pattern = rf'\b{re.escape(skill_lower)}\b'
        return bool(re.search(pattern, text_lower))
    
    def _extract_work_experience(self, sections: Dict[str, str], text: str) -> List[Dict[str, Any]]:
        """Extract work experience entries"""
        experience_text = sections.get('experience', '')
        if not experience_text:
            return []
        
        experiences = []
        
        # Split experience section into individual entries
        # Look for company/position patterns
        entries = self._split_experience_entries(experience_text)
        
        for entry in entries:
            exp_data = self._parse_experience_entry(entry)
            if exp_data:
                experiences.append(exp_data)
        
        return experiences
    
    def _split_experience_entries(self, text: str) -> List[str]:
        """Split experience section into individual job entries"""
        entries = []
        lines = text.split('\n')
        current_entry = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this looks like a new job entry (has dates and company/position info)
            if self._looks_like_job_header(line) and current_entry:
                entries.append('\n'.join(current_entry))
                current_entry = [line]
            else:
                current_entry.append(line)
        
        # Add the last entry
        if current_entry:
            entries.append('\n'.join(current_entry))
        
        return entries
    
    def _looks_like_job_header(self, line: str) -> bool:
        """Check if a line looks like a job header"""
        # Look for date patterns and job-like keywords
        has_dates = bool(re.search(r'\d{4}', line))
        has_job_keywords = any(keyword in line.lower() for keyword in 
                              ['engineer', 'developer', 'manager', 'analyst', 'specialist', 'coordinator'])
        
        return has_dates or has_job_keywords
    
    def _parse_experience_entry(self, entry_text: str) -> Optional[Dict[str, Any]]:
        """Parse individual work experience entry"""
        lines = entry_text.strip().split('\n')
        if not lines:
            return None
        
        # Extract basic information
        company = ""
        job_title = ""
        duration = ""
        responsibilities = []
        technologies = []
        
        # Parse first few lines for company, title, and dates
        for i, line in enumerate(lines[:3]):
            line = line.strip()
            
            # Try to extract company and job title
            if i == 0:
                # First line often contains job title or company
                if any(keyword in line.lower() for keyword in ['engineer', 'developer', 'manager']):
                    job_title = line
                else:
                    company = line
            elif i == 1:
                if not job_title and any(keyword in line.lower() for keyword in ['engineer', 'developer', 'manager']):
                    job_title = line
                elif not company:
                    company = line
            
            # Look for date patterns
            date_match = re.search(r'(\d{4}\s*[-–]\s*\d{4}|\d{4}\s*[-–]\s*present|\w+\s+\d{4}\s*[-–]\s*\w+\s+\d{4})', line, re.IGNORECASE)
            if date_match:
                duration = date_match.group(1)
        
        # Extract responsibilities (bullet points or numbered items)
        for line in lines:
            line = line.strip()
            if line.startswith(('•', '-', '*', '◦')) or re.match(r'^\d+\.', line):
                responsibility = re.sub(r'^[•\-*◦\d.]+\s*', '', line)
                if len(responsibility) > 10:
                    responsibilities.append(responsibility)
        
        # Extract technologies mentioned in the entry
        for category_skills in self.technical_skills_db.values():
            for skill in category_skills:
                if self._skill_mentioned(entry_text, skill):
                    technologies.append(skill)
        
        # Calculate duration in months (approximate)
        duration_months = self._parse_duration_to_months(duration)
        
        if job_title or company:
            return {
                'company': company,
                'job_title': job_title,
                'duration': duration,
                'duration_months': duration_months,
                'responsibilities': responsibilities,
                'technologies_used': list(set(technologies))
            }
        
        return None
    
    def _parse_duration_to_months(self, duration_str: str) -> int:
        """Convert duration string to approximate months"""
        if not duration_str:
            return 0
        
        # Extract years from duration
        years_match = re.findall(r'(\d{4})', duration_str)
        if len(years_match) >= 2:
            start_year = int(years_match[0])
            end_year = int(years_match[1]) if years_match[1] else 2024
            return (end_year - start_year) * 12
        
        return 0
    
    def _extract_projects(self, sections: Dict[str, str], text: str) -> List[Dict[str, Any]]:
        """Extract project information"""
        projects_text = sections.get('projects', '')
        if not projects_text:
            return []
        
        projects = []
        entries = self._split_project_entries(projects_text)
        
        for entry in entries:
            project_data = self._parse_project_entry(entry)
            if project_data:
                projects.append(project_data)
        
        return projects
    
    def _split_project_entries(self, text: str) -> List[str]:
        """Split projects section into individual project entries"""
        entries = []
        lines = text.split('\n')
        current_entry = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this looks like a new project (project name or bullet point)
            if (line.startswith(('•', '-', '*')) or 
                (not line.startswith(' ') and len(line) < 100 and current_entry)):
                
                if current_entry:
                    entries.append('\n'.join(current_entry))
                    current_entry = [line]
                else:
                    current_entry.append(line)
            else:
                current_entry.append(line)
        
        # Add the last entry
        if current_entry:
            entries.append('\n'.join(current_entry))
        
        return entries
    
    def _parse_project_entry(self, entry_text: str) -> Optional[Dict[str, Any]]:
        """Parse individual project entry"""
        lines = entry_text.strip().split('\n')
        if not lines:
            return None
        
        project_name = ""
        description = ""
        technologies = []
        github_link = ""
        demo_link = ""
        
        # Extract project name (usually first line)
        first_line = lines[0].strip()
        project_name = re.sub(r'^[•\-*◦\d.]+\s*', '', first_line)
        
        # Extract description (combine remaining lines)
        if len(lines) > 1:
            description = ' '.join(line.strip() for line in lines[1:] if line.strip())
        
        # Extract GitHub links
        github_match = re.search(r'(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9_/-]+', entry_text, re.IGNORECASE)
        if github_match:
            github_link = github_match.group(0)
        
        # Extract demo/live links
        demo_pattern = r'(?:demo|live|deployed).*?(?:https?://[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?:/[^\s]*)?)'
        demo_match = re.search(demo_pattern, entry_text, re.IGNORECASE)
        if demo_match:
            url_match = re.search(r'https?://[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?:/[^\s]*)?', demo_match.group(0))
            if url_match:
                demo_link = url_match.group(0)
        
        # Extract technologies
        for category_skills in self.technical_skills_db.values():
            for skill in category_skills:
                if self._skill_mentioned(entry_text, skill):
                    technologies.append(skill)
        
        if project_name:
            return {
                'name': project_name,
                'description': description,
                'technologies_used': list(set(technologies)),
                'github_url': github_link,
                'demo_url': demo_link,
                'role': 'Developer'  # Default role, could be enhanced
            }
        
        return None
    
    def _extract_education(self, sections: Dict[str, str], text: str) -> List[Dict[str, Any]]:
        """Extract education information"""
        education_text = sections.get('education', '')
        if not education_text:
            return []
        
        education_entries = []
        lines = education_text.split('\n')
        
        current_entry = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Look for degree patterns
            degree_match = re.search(r'(bachelor|master|phd|doctorate|associate|diploma|certificate).*?(computer science|engineering|mathematics|physics|business)', line, re.IGNORECASE)
            if degree_match:
                if current_entry:
                    education_entries.append(current_entry)
                
                current_entry = {
                    'degree': degree_match.group(0),
                    'institution': '',
                    'graduation_year': '',
                    'gpa': ''
                }
            
            # Look for institution names (universities, colleges)
            if any(keyword in line.lower() for keyword in ['university', 'college', 'institute', 'school']):
                if current_entry and not current_entry.get('institution'):
                    current_entry['institution'] = line
            
            # Look for graduation years
            year_match = re.search(r'\b(19|20)\d{2}\b', line)
            if year_match and current_entry:
                current_entry['graduation_year'] = year_match.group(0)
            
            # Look for GPA
            gpa_match = re.search(r'gpa:?\s*(\d+\.?\d*)', line, re.IGNORECASE)
            if gpa_match and current_entry:
                current_entry['gpa'] = gpa_match.group(1)
        
        # Add the last entry
        if current_entry:
            education_entries.append(current_entry)
        
        return education_entries
    
    def _extract_certifications(self, sections: Dict[str, str], text: str) -> List[str]:
        """Extract certifications"""
        cert_text = sections.get('certifications', '')
        if not cert_text:
            return []
        
        certifications = []
        lines = cert_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if line and len(line) > 3:
                # Clean up bullet points
                cert = re.sub(r'^[•\-*◦\d.]+\s*', '', line)
                if cert:
                    certifications.append(cert)
        
        return certifications
    
    def _extract_languages(self, sections: Dict[str, str], text: str) -> List[str]:
        """Extract spoken languages (not programming languages)"""
        lang_text = sections.get('languages', '')
        if not lang_text:
            return []
        
        # Common spoken languages
        spoken_languages = [
            'English', 'Spanish', 'French', 'German', 'Italian', 'Portuguese',
            'Chinese', 'Japanese', 'Korean', 'Arabic', 'Hindi', 'Russian',
            'Dutch', 'Swedish', 'Norwegian', 'Danish', 'Finnish', 'Polish',
        ]
        
        found_languages = []
        for language in spoken_languages:
            if self._skill_mentioned(lang_text, language):
                found_languages.append(language)
        
        return found_languages
    
    def _extract_achievements(self, sections: Dict[str, str], text: str) -> List[str]:
        """Extract achievements and awards"""
        achievements_text = sections.get('achievements', '')
        if not achievements_text:
            return []
        
        achievements = []
        lines = achievements_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if line and len(line) > 5:
                achievement = re.sub(r'^[•\-*◦\d.]+\s*', '', line)
                if achievement:
                    achievements.append(achievement)
        
        return achievements
    
    def _extract_publications(self, sections: Dict[str, str], text: str) -> List[str]:
        """Extract publications"""
        pub_text = sections.get('publications', '')
        if not pub_text:
            return []
        
        publications = []
        lines = pub_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if line and len(line) > 10:
                publication = re.sub(r'^[•\-*◦\d.]+\s*', '', line)
                if publication:
                    publications.append(publication)
        
        return publications
    
    def _extract_volunteer_experience(self, sections: Dict[str, str], text: str) -> List[Dict[str, Any]]:
        """Extract volunteer experience"""
        volunteer_text = sections.get('volunteer', '')
        if not volunteer_text:
            return []
        
        volunteer_entries = []
        entries = self._split_experience_entries(volunteer_text)
        
        for entry in entries:
            vol_data = self._parse_volunteer_entry(entry)
            if vol_data:
                volunteer_entries.append(vol_data)
        
        return volunteer_entries
    
    def _parse_volunteer_entry(self, entry_text: str) -> Optional[Dict[str, Any]]:
        """Parse individual volunteer experience entry"""
        lines = entry_text.strip().split('\n')
        if not lines:
            return None
        
        organization = ""
        role = ""
        duration = ""
        description = ""
        
        # Extract organization and role from first few lines
        for i, line in enumerate(lines[:2]):
            line = line.strip()
            if i == 0:
                organization = line
            elif i == 1:
                role = line
        
        # Look for duration
        for line in lines:
            date_match = re.search(r'(\d{4}\s*[-–]\s*\d{4}|\d{4}\s*[-–]\s*present)', line, re.IGNORECASE)
            if date_match:
                duration = date_match.group(1)
                break
        
        # Extract description
        desc_lines = []
        for line in lines:
            line = line.strip()
            if line.startswith(('•', '-', '*')):
                desc_lines.append(re.sub(r'^[•\-*◦\d.]+\s*', '', line))
        
        description = ' '.join(desc_lines) if desc_lines else ''
        
        if organization:
            return {
                'organization': organization,
                'role': role,
                'duration': duration,
                'description': description
            }
        
        return None
    
    def _calculate_confidence(self, result: ExtractedResumeData) -> float:
        """Calculate extraction confidence score"""
        confidence_factors = []
        
        # Contact information completeness (40% weight)
        contact_score = 0
        if result.full_name:
            contact_score += 0.3
        if result.email:
            contact_score += 0.4
        if result.phone:
            contact_score += 0.2
        if result.linkedin_url or result.github_url:
            contact_score += 0.1
        
        confidence_factors.append(('contact', contact_score, 0.4))
        
        # Skills extraction (25% weight)
        total_skills = len(result.technical_skills) + len(result.soft_skills)
        skills_score = min(1.0, total_skills / 15)  # Normalize to 15 skills
        confidence_factors.append(('skills', skills_score, 0.25))
        
        # Experience extraction (20% weight)
        exp_score = min(1.0, len(result.work_experience) / 3)  # Normalize to 3 jobs
        confidence_factors.append(('experience', exp_score, 0.2))
        
        # Section detection (10% weight)
        section_score = min(1.0, len(result.sections_found) / 6)  # Normalize to 6 sections
        confidence_factors.append(('sections', section_score, 0.1))
        
        # Education extraction (5% weight)
        edu_score = min(1.0, len(result.education))
        confidence_factors.append(('education', edu_score, 0.05))
        
        # Calculate weighted average
        total_confidence = sum(score * weight for _, score, weight in confidence_factors)
        
        return round(total_confidence, 3)